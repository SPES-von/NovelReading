# -*- coding: utf-8 -*-
"""生成《形势与政策VIII》结课答卷 Word 文档（五号宋体、单倍行距、文末编号注释）。"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt
from docx.oxml.ns import qn


def set_run_east_asia_font(run, name="宋体", size_pt=10.5, bold=False):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    r = run._element.rPr
    if r is None:
        run._element.get_or_add_rPr()
        r = run._element.rPr
    rFonts = r.rFonts
    if rFonts is None:
        from docx.oxml import OxmlElement

        rFonts = OxmlElement("w:rFonts")
        r.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)


def format_paragraph(paragraph, first_line_indent_cm=0.74):
    from docx.shared import Cm

    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.first_line_indent = Cm(first_line_indent_cm)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run(text)
    set_run_east_asia_font(run, "宋体", 10.5, False)
    return p


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(12)
    run = p.add_run(text)
    set_run_east_asia_font(run, "宋体", 15, False)  # 小三左右作标题层次，仍用宋体不加粗
    return p


def add_footnote_heading(doc):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run("脚注")
    set_run_east_asia_font(run, "宋体", 10.5, False)
    return p


def add_footnote_item(doc, marker, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.left_indent = Pt(21)
    p.paragraph_format.first_line_indent = Pt(-21)
    run = p.add_run(f"{marker}{text}")
    set_run_east_asia_font(run, "宋体", 10.5, False)
    return p


def main():
    out = Path(__file__).resolve().parent.parent / "请填写学号姓名《形势与政策VIII》结课考试答卷.docx"

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(10.5)
    if normal._element.rPr is None:
        normal._element.get_or_add_rPr()
    rpr = normal._element.rPr
    if rpr.rFonts is None:
        from docx.oxml import OxmlElement

        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rpr.rFonts.set(qn("w:eastAsia"), "宋体")

    add_title(
        doc,
        "数字技术赋能文化遗产“活起来”与“传下去”\n——兼论VR、AR与大数据的实践路径与个人思考",
    )

    paras = [
        "党的十八大以来，我国文化遗产工作在总体要求上突出“保护第一、加强管理、挖掘价值、有效利用、让文物活起来”。创造性转化与创新性发展（“两创”）进一步提示：遗产保护不是把历史封存在库房，而要在当代生活中重新生成意义。与此同时，气候风险、旅游强度、城市更新与突发事件等因素，使许多遗产地长期处于高负荷状态；若仅依赖传统修缮与有限的线下开放，很难在“安全可控”与“人民共享”之间建立可持续机制。正因如此，国家将文化数字化纳入顶层设计，强调用数字基础设施提升文化资源的保存、再现与服务能力，这预示着数字化正从零星试点走向系统性能力建设。",
        "在此语境下讨论“活起来”，我更愿意把它理解为一种公共文化过程：公众不仅得以观看，更能在理解、讨论与再阐释的链条中获得主体性。虚拟现实（VR）擅长营造可进入的空间现场，使人以身体尺度去感知殿堂高度、甬道节奏与壁画设色；增强现实（AR）则能把数字图层精确地对齐在真实器物或遗址表面，让无法直接观察的内层信息、已消失的装饰元素以“对照阅读”的方式出现。它们共同的作用，是把考古地层关系、材料老化机理与历史叙事转化为可感可知的内容，从而降低专业门槛，扩大参与人口。",
        "从教育传播角度看，数字展陈还能突破馆舍边界：偏远地区学校通过线上资源获得与中心城市相近的某些课程素材，进城务工人员子女也能在移动终端上接触到高质量的文物故事。当然，这类公平并不自动降临——如果内容分发过度依赖昂贵设备与高速网络，反而可能放大差距。因而，“活起来”同时要求我们推动终端多样化、低带宽版本与公共服务平台的可达性，这才更符合共同富裕与社会公平的价值取向。",
        "就个人学习经历而言，我更在意体验设计中的三条底线。第一是证据分层：凡是进入公众界面的信息，都应对“考古实证”“文献互证”“学者推断”“艺术想象”作清晰标注，避免把复原动画当作无可置疑的史实。第二是问题导向：与其堆砌百科词条，不如设置可反思的任务，让观众在虚拟情境里理解“为何限制闪光灯”“为何某些洞窟需要轮流开放”。第三是线上线下互证：数字产品应当激发人到现场去核对细节的欲望，而不是以便利之名消解对物质真实的尊重。否则，“热度”来得快，认同却很难沉淀。",
        "与“活起来”并行的，是“传下去”的长期维度：把时间轴拉长，文化遗产真正稀缺的是不可逆的信息损失。高精度影像、三维激光扫描、多光谱成像等采集手段，使石窟、壁画、古籍、木构建筑等对象能够在形态、材料与微观损伤等层面留下可比对、可追溯的“数字底本”。当原件因自然老化或突发事件面临风险时，数字档案至少能提供修复与研究的历史参照；在日常管理中，规范化的命名、元数据与备份校验，则是防止“硬盘损坏就等于遗产信息灭失”的基础工程。对我而言，数字化首先是一项严肃的档案制度，其次才是视觉奇观。",
        "把预防性保护再推进一步，各类传感监测、环境数据与巡检记录如果进入同一数据治理框架，就会形成面向风险的预警能力。这也意味着文博机构需要跨部门协同：保护、物业、信息化与科研人员共享标准接口，而不是各建一套互不连通的系统。数字化的深层收益，是把碎片化的看管经验转化为可积累、可复核的证据链，使“发现问题—研判原因—采取措施”的闭环更可操作。",
        "大数据在这一链条上扮演“连接器”。观众在展线上的驻留与跳转，可以帮助策展者发现叙事断点；跨区域风格特征库与知识图谱，则有助于研究古代工匠流动与材料贸易。治理层面，资源目录与统计模型也能支撑资金与人才的更精准投放。但数据规模越大，越要把边界划清：哪些数据可以开放，哪些必须在专网内使用；如何对个人信息脱敏；如何避免算法推荐把遗产观看简化为猎奇消费。必须在法治框架内推进项目，才可能把“数治”真正变成善治。",
        "近年来，人工智能在图像检索、碑文识读与标注辅助方面展现出潜力，客观上能提升研究效率。但我坚持一点：生成式人工智能若用于“补足缺失的历史情节”，必须以显著提示标明其虚构性质，并由学者审校后才能进入公共教育产品，否则会混淆公众的历史感。技术越是强大，阐释责任越要明确，这是我在课堂讨论中反复申说的立场。",
        "站在青年学生的位置，我时常提醒自己：我们既是数字内容的受众，也可能在未来进入公共文化服务、文化创意或技术研发的相关岗位。越是身处“会用工具”的一代，越要在日常学习中养成查证来源、区分事实与推断的习惯，把对前沿技术的兴奋感与对方法边界的克制结合起来。否则，真诚的文化热情也可能在无意中放大误读，背离遗产保护的初衷。",
        "当“保存—研究—传播—反馈”形成闭环，数字化才算完成从工具到机制的跃迁。线上提问与线下参观可以互相校正，监测异常可以反向触发保护行动，公众参与也能为地方叙事提供细节补充。由此，文化遗产体系在城市风险、气候波动与技术变迁面前，会表现出更强的韧性与敏捷性，而不是被动等待危机发生后再补救。",
        "也必须承认，技术应用会带来新的伦理张力：娱乐化过度可能削弱遗产的庄重感；平台算法可能形塑刻板印象；数字复制品一旦被误读为“同等于原件”，会动摇人们对物质载体与工匠传统的敬畏。因此，我倾向于在项目立项阶段就引入风险评估与社区协商，把传承人、当地居民与多元群体纳入共创，让数字阐释成为对话的平台，而不是单向的技术炫示。业界的持续讨论提醒我们：传播效果与学术严谨之间需要动态平衡，任何“唯流量”取向都值得警惕。",
        "展望未来，我国文化遗产数字化若要从“会做”走向“做久”，关键在于制度化：统一元数据标准，明确长期保存责任单位，建立可持续运维与审计机制，把短期展演沉淀为可继承的数字资产。唯有如此，数字技术才能真正使文化遗产在当下有感召力，在未来仍可验证、可学习、可对话，实现既“活起来”，也“传下去”。",
    ]

    markers = ["¹", "²", "³", "⁴", "⁵"]
    for i, t in enumerate(paras):
        if i == 0:
            t = t + markers[0]
        elif i == 1:
            t = t + markers[1]
        elif i == 4:
            t = t + markers[2]
        elif i == 6:
            t = t + markers[3]
        elif i == 10:
            t = t + markers[4]
        add_body_paragraph(doc, t)

    footnotes = [
        (
            "¹",
            "中共中央办公厅、国务院办公厅：《关于推进实施国家文化数字化战略的意见》，对文化数字化总体方向与重点任务作出系统部署，可结合公开文本查阅。",
        ),
        (
            "²",
            "联合国教科文组织（UNESCO）在遗产保护与公众参与相关文件中，对数字化阐释与包容性传播有持续讨论，可作为国际参照。",
        ),
        (
            "³",
            "国家文物局及业内技术规范、公开案例中关于文物数字化采集与档案建设的实践，可为“数字底本”思路提供行业语境。",
        ),
        (
            "⁴",
            "《中华人民共和国数据安全法》《中华人民共和国个人信息保护法》等法律框架，对数据处理活动提出基本要求，数字文博项目应与之对齐。",
        ),
        (
            "⁵",
            "参见《中国文物报》等专业媒体关于数字化展览、虚拟体验伦理边界的相关评论与讨论。",
        ),
    ]

    add_footnote_heading(doc)
    for m, ft in footnotes:
        add_footnote_item(doc, m, ft)

    doc.save(out)
    print("已生成:", out)

    # 统计正文汉字约数（不含脚注区标题与脚注条文内数字符号的粗略统计）
    body = "".join(paras)
    import re

    han = re.findall(r"[\u4e00-\u9fff]", body)
    print("正文汉字约:", len(han))


if __name__ == "__main__":
    main()
